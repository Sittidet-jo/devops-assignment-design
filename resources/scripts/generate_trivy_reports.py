#!/usr/bin/env python3
"""
Trivy Vulnerability Report Generator
PDF + DOCX + XLSX from trivy JSON report
Bilingual: Thai + English
"""
import json, sys, os
from datetime import datetime
from collections import defaultdict

from report_utils import (
    load_font, make_pdf_styles, pdf_colors, pdf_base_tbl_style,
    make_pdf_table, pdf_meta_table, pdf_gate_paragraph,
    docx_make_table, docx_add_heading, docx_gate_paragraph,docx_color_cell_text, docx_set_default_font,
    xlsx_styles, xlsx_hcell, xlsx_dcell, xlsx_title_row, xlsx_section_header,
)

def sev_color(sev: str) -> str:
    return {"CRITICAL": "#C0392B", "HIGH": "#E67E22", "MEDIUM": "#D4AC0D",
            "LOW": "#2E86C1", "UNKNOWN": "#7F8C8D"}.get(sev.upper(), "#000000")

def sev_order(sev: str) -> int:
    return {"CRITICAL": 0, "HIGH": 1, "MEDIUM": 2, "LOW": 3, "UNKNOWN": 4}.get((sev or "").upper(), 5)

def parse_criteria(criteria_str: str) -> dict:
    import re
    crit = re.search(r'Critical\s*[<=]+\s*(\d+)', criteria_str or '', re.IGNORECASE)
    high = re.search(r'High\s*[<=]+\s*(\d+)',    criteria_str or '', re.IGNORECASE)
    return {
        "critical_threshold": int(crit.group(1)) if crit else 0,
        "high_threshold":     int(high.group(1)) if high else 5,
    }

def evaluate_gate(counts: dict, thresholds: dict) -> str:
    c = int(counts.get("CRITICAL", 0)); h = int(counts.get("HIGH", 0))
    if c > thresholds["critical_threshold"] or h > thresholds["high_threshold"]: return "FAIL"
    if c > 0 or h > 0: return "WARN"
    return "PASS"

def parse_report(report: dict) -> dict:
    image_name = report.get("ArtifactName", "unknown")
    meta       = report.get("Metadata", {}); os_info = meta.get("OS", {})
    os_str     = f"{os_info.get('Family', '')} {os_info.get('Name', '')}".strip()
    counts     = defaultdict(int); all_vulns = []

    for result in report.get("Results", []):
        target = result.get("Target", ""); pkg_type = result.get("Type", "")
        for v in (result.get("Vulnerabilities") or []):
            sev = (v.get("Severity") or "UNKNOWN").upper(); counts[sev] += 1
            cvss_score = ""
            for src in ["nvd", "ghsa", "redhat"]:
                cvss_data = v.get("CVSS", {})
                if src in cvss_data and cvss_data[src].get("V3Score"):
                    cvss_score = f"{cvss_data[src]['V3Score']:.1f}"; break
            all_vulns.append({
                "id":        v.get("VulnerabilityID", ""), "pkg":       v.get("PkgName", ""),
                "installed": v.get("InstalledVersion", ""), "fixed":     v.get("FixedVersion", "") or "No fix",
                "severity":  sev, "title": v.get("Title", v.get("VulnerabilityID", "")),
                "cvss":      cvss_score, "target": target, "type": pkg_type,
            })
    all_vulns.sort(key=lambda x: (sev_order(x["severity"]), x["pkg"]))
    return {
        "image_name": image_name, "os": os_str,
        "scanned_at": datetime.now().strftime("%Y-%m-%d %H:%M"),
        "counts":     dict(counts), "total": sum(counts.values()),
        "crit_high":  [v for v in all_vulns if v["severity"] in ("CRITICAL", "HIGH")],
        "medium":     [v for v in all_vulns if v["severity"] == "MEDIUM"],
        "all_vulns":  all_vulns,
    }

TH = {
    # Full-width rows — bilingual
    "title":     "รายงานผลการสแกนช่องโหว่ / Trivy Vulnerability Scan Report",
    "summary":   "สรุปผลการสแกน / Scan Summary",
    "crit_high": "Critical & High — ทั้งหมด / All",
    "medium":    "Medium — ทั้งหมด / All",
    "no_vuln":   "ไม่พบช่องโหว่ / No vulnerabilities found",
    "gate_pass": "✅ PASS — ไม่พบ Critical/High / No Critical or High Vulnerabilities",
    "gate_warn": "⚠️ WARN — พบ Critical/High แต่อยู่ในเกณฑ์ / Within Accepted Thresholds",
    "gate_fail": "❌ FAIL — พบ Critical/High เกินเกณฑ์ (Pipeline: UNSTABLE) / Threshold Exceeded",
    "scanned":   "สแกนเวลา / Scanned At",
    "threshold": "เกณฑ์ / Threshold",
    "repo":      "GitLab Repository",
    "branch":    "Branch",
    "repo_url":  "Repository URL",
    "image":     "Image",
    "os":        "OS / Base Image",
    "severity":  "Severity",
    "count":     "Count",
    "total":     "Total",
    "cve_id":    "CVE ID",
    "package":   "Package",
    "installed": "Installed",
    "fixed":     "Fixed",
    "cvss":      "CVSS",
    "title_col": "Title",
}

GATE_COLOR = {"PASS": "#1E8449", "WARN": "#E67E22", "FAIL": "#C0392B"}
GATE_TEXT  = {k: TH[f"gate_{k.lower()}"] for k in GATE_COLOR}

def generate_pdf(data: dict, out_path: str) -> None:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.lib.units import mm
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable, PageBreak)

    FN, FB = load_font()
    ST = make_pdf_styles(FN, FB); C = pdf_colors()
    sT, sH1, sN, sS, sB, sHDR = ST["sT"], ST["sH1"], ST["sN"], ST["sS"], ST["sB"], ST["sHDR"]

    W, _ = A4; CW = W - 30 * mm
    doc = SimpleDocTemplate(out_path, pagesize=A4, leftMargin=15*mm, rightMargin=15*mm, topMargin=18*mm, bottomMargin=18*mm)

    counts = data["counts"]; story = []

    story.append(Paragraph(TH["title"], sT))
    story.append(HRFlowable(width="100%", thickness=2, color=C["BLUE"]))
    story.append(Spacer(1, 3*mm))

    repo     = os.environ.get("REPO_PATH", "-")
    branch   = os.environ.get("BRANCH_NAME", "-")
    repo_url = os.environ.get("REPO_URL", "-")

    kv = [
        (TH["scanned"], data["scanned_at"]), 
        (TH["image"], data["image_name"]), 
        (TH["os"], data["os"] or "-"),
        (TH["repo"], repo),
        (TH["branch"], branch),
        (TH["repo_url"], repo_url)
    ]
    if data.get("criteria"):
        kv.append((TH["threshold"], data["criteria"]))
    story.append(pdf_meta_table(kv, [40*mm, CW - 40*mm], C, ST))
    story.append(Spacer(1, 3*mm))

    gate = data.get("gate", "PASS")
    story.append(pdf_gate_paragraph(GATE_TEXT.get(gate, ""), GATE_COLOR.get(gate, "#7F8C8D"), FB))
    story.append(Spacer(1, 4*mm))

    story.append(Paragraph(TH["summary"], sH1))
    sev_rows = []; sev_idx = {}
    for sev in ["CRITICAL", "HIGH", "MEDIUM", "LOW", "UNKNOWN"]:
        cnt = counts.get(sev, 0)
        if cnt > 0: sev_idx[len(sev_rows)] = sev; sev_rows.append((sev, str(cnt)))
    sev_rows.append((TH["total"], str(data["total"])))
    sum_style = pdf_base_tbl_style(len(sev_rows), C)
    for i, sev in sev_idx.items():
        c = colors.HexColor(sev_color(sev))
        sum_style += [("TEXTCOLOR", (0, i+1), (0, i+1), c), ("FONTNAME", (0, i+1), (0, i+1), FB)]
    sum_style += [("FONTNAME",   (0, len(sev_rows)), (1, len(sev_rows)), FB),
                  ("BACKGROUND", (0, len(sev_rows)), (-1, len(sev_rows)), C["LBLU"])]
    st = Table(
        [[Paragraph(TH["severity"], sHDR), Paragraph(TH["count"], sHDR)]] +
        [[Paragraph(f'<font color="{sev_color(s)}"><b>{s}</b></font>', sS), Paragraph(c, sS)] for s, c in sev_rows],
        colWidths=[100*mm, CW - 100*mm],
    )
    st.setStyle(TableStyle(sum_style)); story.append(st); story.append(Spacer(1, 4*mm))

    def vuln_table(vuln_list: list, label: str) -> None:
        story.append(Paragraph(f"<b>{label}</b> ({len(vuln_list)} items)", sN))
        story.append(Spacer(1, 1*mm))
        if not vuln_list:
            story.append(Paragraph(TH["no_vuln"], sS)); story.append(Spacer(1, 3*mm)); return
        hdrs = [TH["cve_id"], TH["severity"], TH["package"], TH["installed"], TH["fixed"], TH["cvss"], TH["title_col"]]
        cws = [26*mm, 16*mm, 22*mm, 22*mm, 22*mm, 12*mm, CW - 120*mm]
        v_style = pdf_base_tbl_style(len(vuln_list), C)
        for i, v in enumerate(vuln_list):
            c = colors.HexColor(sev_color(v["severity"]))
            v_style += [("TEXTCOLOR", (1, i+1), (1, i+1), c), ("FONTNAME", (1, i+1), (1, i+1), FB)]
        rows = [[v["id"], f'<font color="{sev_color(v["severity"])}"><b>{v["severity"]}</b></font>', 
                 v["pkg"], v["installed"], v["fixed"], v["cvss"],
                 v["title"][:65] + "…" if len(v["title"]) > 65 else v["title"]] for v in vuln_list]
        t = Table(
            [[Paragraph(h, sHDR) for h in hdrs]] +
            [[Paragraph(str(c), sS) for c in row] for row in rows],
            colWidths=cws,
        )
        t.setStyle(TableStyle(v_style)); story.append(t); story.append(Spacer(1, 4*mm))

    story.append(PageBreak())
    vuln_table(data["crit_high"], TH["crit_high"])
    story.append(PageBreak())
    vuln_table(data["medium"],    TH["medium"])
    doc.build(story)

def generate_docx(data: dict, out_path: str) -> None:
    from docx import Document
    from docx.shared import RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    TW = 17.0
    doc = Document()
    docx_set_default_font(doc, "Arial")
    for sec in doc.sections:
        sec.top_margin = sec.bottom_margin = Cm(2)
        sec.left_margin = sec.right_margin = Cm(2)

    counts = data["counts"]
    t = doc.add_heading(TH["title"], 0); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in t.runs: r.font.color.rgb = RGBColor.from_string("1A5276")

    repo     = os.environ.get("REPO_PATH", "-")
    branch   = os.environ.get("BRANCH_NAME", "-")
    repo_url = os.environ.get("REPO_URL", "-")

    meta_rows = [
        [TH["scanned"], data["scanned_at"]], 
        [TH["image"], data["image_name"]], 
        [TH["os"], data["os"] or "-"],
        [TH["repo"], repo],
        [TH["branch"], branch],
        [TH["repo_url"], repo_url]
    ]
    if data.get("criteria"): meta_rows.append([TH["threshold"], data["criteria"]])
    docx_make_table(doc, [], meta_rows, [4.5, TW - 4.5])
    doc.add_paragraph()

    gate = data.get("gate", "PASS")
    docx_gate_paragraph(doc, GATE_TEXT.get(gate, ""), GATE_COLOR.get(gate, "#7F8C8D"))
    doc.add_paragraph()

    docx_add_heading(doc, TH["summary"])
    sev_rows = [[sev, str(counts.get(sev, 0))] for sev in ["CRITICAL", "HIGH", "MEDIUM", "LOW", "UNKNOWN"] if counts.get(sev, 0) > 0]
    sev_rows.append([TH["total"], str(data["total"])])
    t_sum = docx_make_table(doc, [TH["severity"], TH["count"]], sev_rows, [8.0, TW - 8.0])
    for i, row in enumerate(sev_rows[:-1]):
        docx_color_cell_text(t_sum, i + 1, 0, sev_color(row[0]))
    doc.add_paragraph()

    vuln_headers = [TH["cve_id"], TH["severity"], TH["package"], TH["installed"], TH["fixed"], TH["cvss"], TH["title_col"]]
    vuln_widths  = [3.0, 2.0, 2.8, 2.2, 2.2, 1.2, TW - 13.4]

    def add_vuln_table(vuln_list: list, label: str) -> None:
        docx_add_heading(doc, label, level=2)
        if not vuln_list: doc.add_paragraph(TH["no_vuln"]); return
        t_v = docx_make_table(doc, vuln_headers,
                        [[v["id"], v["severity"], v["pkg"], v["installed"],
                          v["fixed"], v["cvss"], v["title"][:100]] for v in vuln_list],
                        vuln_widths)
        for i, v in enumerate(vuln_list):
            docx_color_cell_text(t_v, i + 1, 1, sev_color(v["severity"]))
        doc.add_paragraph()

    doc.add_page_break()
    add_vuln_table(data["crit_high"], TH["crit_high"])
    doc.add_page_break()
    add_vuln_table(data["medium"],    TH["medium"])
    doc.save(out_path)

def generate_xlsx(data: dict, out_path: str) -> None:
    from openpyxl import Workbook
    from openpyxl.styles import Font

    wb = Workbook(); counts = data["counts"]
    S  = xlsx_styles()

    ws1 = wb.active; ws1.title = "📊 Summary"
    xlsx_title_row(ws1, TH["title"], S, end_col=3)
    
    repo     = os.environ.get("REPO_PATH", "-")
    branch   = os.environ.get("BRANCH_NAME", "-")
    repo_url = os.environ.get("REPO_URL", "-")

    _meta = [
        (TH["scanned"], data["scanned_at"]), 
        (TH["image"], data["image_name"]), 
        (TH["os"], data["os"] or "-"),
        (TH["repo"], repo),
        (TH["branch"], branch),
        (TH["repo_url"], repo_url)
    ]
    for i, (k, v) in enumerate(_meta, start=2):
        c1 = ws1.cell(row=i, column=1, value=k); c1.font = S["BBLUE"]; c1.fill = S["LFILL"]; c1.border = S["BDR"]; c1.alignment = S["LFT"]
        ws1.merge_cells(f"B{i}:C{i}"); c2 = ws1.cell(row=i, column=2, value=v); c2.font = S["NRM"]; c2.border = S["BDR"]; c2.alignment = S["LFT"]

    row = 2 + len(_meta) + 1; gate = data.get("gate", "PASS"); gate_hex = GATE_COLOR.get(gate, "7F8C8D").lstrip("#")
    ws1.merge_cells(f"A{row}:C{row}"); gc = ws1.cell(row=row, column=1, value=GATE_TEXT.get(gate, ""))
    gc.font = Font(bold=True, color=gate_hex, name="Arial", size=11)
    gc.alignment = S["CTR"]; gc.border = S["BDR"]; gc.fill = S["WFILL"]
    ws1.row_dimensions[row].height = 22; row += 2

    xlsx_section_header(ws1, row, TH["summary"], S, end_col=3); row += 1
    xlsx_hcell(ws1, row, 1, TH["severity"], S, 20); xlsx_hcell(ws1, row, 2, TH["count"], S, 14); row += 1
    for i, sev in enumerate(["CRITICAL", "HIGH", "MEDIUM", "LOW", "UNKNOWN"]):
        cnt = counts.get(sev, 0)
        if cnt == 0: continue
        alt = i % 2 == 0; sc = sev_color(sev).lstrip("#")
        xlsx_dcell(ws1, row, 1, sev, S, alt, bold=True, color=sc)
        xlsx_dcell(ws1, row, 2, cnt, S, alt); row += 1
    tc = ws1.cell(row=row, column=1, value=TH["total"])
    tc.font = Font(bold=True, name="Arial", size=9); tc.fill = S["LFILL"]; tc.border = S["BDR"]; tc.alignment = S["CTR"]
    vc = ws1.cell(row=row, column=2, value=data["total"])
    vc.font = Font(bold=True, name="Arial", size=9); vc.fill = S["LFILL"]; vc.border = S["BDR"]; vc.alignment = S["CTR"]
    ws1.column_dimensions["A"].width = 28; ws1.column_dimensions["B"].width = 16; ws1.column_dimensions["C"].width = 40

    vuln_headers = [TH["cve_id"], TH["severity"], TH["package"], TH["installed"], TH["fixed"], TH["cvss"], TH["title_col"]]
    vuln_widths  = [22, 13, 20, 16, 16, 8, 55]

    def write_vuln_sheet(ws, vuln_list: list, label: str) -> None:
        xlsx_title_row(ws, label, S, end_col=7)
        for col, (h, w) in enumerate(zip(vuln_headers, vuln_widths), 1):
            xlsx_hcell(ws, 2, col, h, S, w)
        if not vuln_list: ws.cell(row=3, column=1, value=TH["no_vuln"]).font = S["NRM"]; return
        for i, v in enumerate(vuln_list):
            r = i + 3; alt = i % 2 == 0; sc = sev_color(v["severity"]).lstrip("#")
            xlsx_dcell(ws, r, 1, v["id"],        S, alt, align=S["LFT"])
            xlsx_dcell(ws, r, 2, v["severity"],  S, alt, bold=True, color=sc)
            xlsx_dcell(ws, r, 3, v["pkg"],       S, alt, align=S["LFT"])
            xlsx_dcell(ws, r, 4, v["installed"], S, alt, align=S["LFT"])
            xlsx_dcell(ws, r, 5, v["fixed"],     S, alt, align=S["LFT"])
            xlsx_dcell(ws, r, 6, v["cvss"],      S, alt)
            xlsx_dcell(ws, r, 7, v["title"],     S, alt, align=S["LFT"])
            ws.row_dimensions[r].height = 30

    ws2 = wb.create_sheet("🔴 Critical & High")
    write_vuln_sheet(ws2, data["crit_high"], f"{TH['crit_high']} ({len(data['crit_high'])} items)")
    ws3 = wb.create_sheet("🟡 Medium")
    write_vuln_sheet(ws3, data["medium"], f"{TH['medium']} ({len(data['medium'])} items)")
    ws4 = wb.create_sheet("📋 All Vulnerabilities")
    write_vuln_sheet(ws4, data["all_vulns"], f"All Vulnerabilities ({len(data['all_vulns'])} items)")
    wb.save(out_path)

def main():
    json_path = sys.argv[1] if len(sys.argv) > 1 else "trivy-report.json"
    out_dir   = sys.argv[2] if len(sys.argv) > 2 else os.path.dirname(json_path) or "reports"
    criteria  = sys.argv[3] if len(sys.argv) > 3 else None

    if not os.path.exists(json_path):
        print(f"[ERROR] File not found: {json_path}", file=sys.stderr); sys.exit(1)
    os.makedirs(out_dir, exist_ok=True)

    report = json.load(open(json_path, encoding="utf-8"))
    data   = parse_report(report)
    if criteria: data["criteria"] = criteria
    thresholds = parse_criteria(criteria or ""); data["gate"] = evaluate_gate(data["counts"], thresholds)

    project_name = os.environ.get("PROJECT_NAME", "")
    deploy_env   = os.environ.get("DEPLOY_ENV", "")
    date = datetime.now().strftime("%Y%m%d")

    if project_name and deploy_env:
        base = f"trivy-{project_name}-{deploy_env}-{date}"
    else:
        img  = data["image_name"].replace("/", "-").replace(":", "-").split("@")[0]
        base = f"trivy-{img}_{date}"
    print(f"[INFO] Trivy: {data['image_name']}  Gate:{data['gate']}  "
          f"CRITICAL={data['counts'].get('CRITICAL', 0)} HIGH={data['counts'].get('HIGH', 0)}")
    generate_pdf (data, os.path.join(out_dir, f"{base}.pdf"))
    generate_docx(data, os.path.join(out_dir, f"{base}.docx"))
    generate_xlsx(data, os.path.join(out_dir, f"{base}.xlsx"))

if __name__ == "__main__":
    main()