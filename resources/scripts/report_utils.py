#!/usr/bin/env python3
"""
report_utils.py — Shared utilities for all security report generators.

Used by:
  - generate_dast_reports.py      (OWASP ZAP)
  - generate_trivy_reports.py     (Trivy)

Design rules:
  • No report-specific logic here — only generic PDF/DOCX/XLSX helpers.
  • All public functions are pure (no global side-effects).
  • Callers own the Document / SimpleDocTemplate lifecycle.
"""

import os

# ══════════════════════════════════════════════════════════════════════════════
# Font loading
# ══════════════════════════════════════════════════════════════════════════════

def load_font() -> tuple[str, str]:
    """
    Register Sarabun Thai fonts with ReportLab and return (normal, bold) font names.
    Falls back to Helvetica if font files are not found (CI without Thai font support).

    Font files are read from the directory set in the FONT_PATH env var (default /tmp).
    """
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    font_dir = os.environ.get("FONT_PATH", "/tmp")
    try:
        pdfmetrics.registerFont(TTFont("TF",  os.path.join(font_dir, "Sarabun-Regular.ttf")))
        pdfmetrics.registerFont(TTFont("TFB", os.path.join(font_dir, "Sarabun-Bold.ttf")))
        return "TF", "TFB"
    except Exception as e:
        print(f"⚠️  Thai font unavailable ({e}), using Helvetica")
        return "Helvetica", "Helvetica-Bold"


# ══════════════════════════════════════════════════════════════════════════════
# PDF — styles and colours
# ══════════════════════════════════════════════════════════════════════════════

# Brand colour used for all headings and table headers.
_BRAND_BLUE = "#1A5276"


def make_pdf_styles(FN: str, FB: str) -> dict:
    """
    Return a dict of named ParagraphStyles built from the supplied font names.

    Keys: sT (title), sH1, sH2, sN (normal), sS (small), sB (bold-small),
          sNote (grey footnote).
    """
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER

    D = _BRAND_BLUE
    return dict(
        sT    = ParagraphStyle("sT",    fontSize=13,  fontName=FB, alignment=TA_CENTER, spaceAfter=4,  textColor=colors.HexColor(D)),
        sH1   = ParagraphStyle("sH1",   fontSize=11,  fontName=FB, spaceBefore=8,  spaceAfter=3,  textColor=colors.HexColor(D)),
        sH2   = ParagraphStyle("sH2",   fontSize=9.5, fontName=FB, spaceBefore=5,  spaceAfter=2,  textColor=colors.HexColor(D)),
        sN    = ParagraphStyle("sN",    fontSize=8.5, fontName=FN, spaceAfter=2),
        sS    = ParagraphStyle("sS",    fontSize=7.5, fontName=FN),
        sB    = ParagraphStyle("sB",    fontSize=8.5, fontName=FB),
        sHDR  = ParagraphStyle("sHDR",  fontSize=8.5, fontName=FB, leading=15, textColor=colors.white),
        sNote = ParagraphStyle("sNote", fontSize=7.5, fontName=FN, leading=10, spaceAfter=4,  textColor=colors.HexColor("#7F8C8D")),
    )


def pdf_colors() -> dict:
    """
    Return a dict of shared ReportLab Color objects.
    Keys: BLUE, LBLU (light blue), GRD (grid/border), RED (light red), WHT.
    """
    from reportlab.lib import colors

    return dict(
        BLUE = colors.HexColor("#1A5276"),
        LBLU = colors.HexColor("#EBF5FB"),
        GRD  = colors.HexColor("#AED6F1"),
        RED  = colors.HexColor("#FADBD8"),
        WHT  = colors.white,
    )


def pdf_base_tbl_style(n_data_rows: int, C: dict) -> list:
    """
    Return a list of TableStyle commands for a standard alternating-row table.

    n_data_rows: number of data rows (header row is NOT counted).
    C: colour dict from pdf_colors().
    """
    style = [
        ("BACKGROUND",    (0, 0), (-1, 0),  C["BLUE"]),
        ("TEXTCOLOR",     (0, 0), (-1, 0),  C["WHT"]),
        ("GRID",          (0, 0), (-1, -1), 0.4, C["GRD"]),
        ("VALIGN",        (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING",   (0, 0), (-1, -1), 5),
        ("TOPPADDING",    (0, 0), (-1, -1), 8),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
        # Extra padding for header row to prevent Thai bold text from clipping
        ("TOPPADDING",    (0, 0), (-1, 0),  10),
        ("BOTTOMPADDING", (0, 0), (-1, 0),  10),
    ]
    for i in range(n_data_rows):
        if i % 2 == 1:
            style.append(("BACKGROUND", (0, i + 1), (-1, i + 1), C["LBLU"]))
    return style


def make_pdf_table(headers: list, rows: list, col_widths: list,
                   styles: dict, C: dict):
    """
    Build and return a ReportLab Table with standard header/alternating-row styling.

    headers    : list of header strings
    rows       : list of lists of cell content (strings)
    col_widths : list of column widths in ReportLab units
    styles     : dict from make_pdf_styles()
    C          : dict from pdf_colors()
    """
    from reportlab.platypus import Table, TableStyle
    from reportlab.platypus import Paragraph

    sB, sS, sHDR = styles["sB"], styles["sS"], styles["sHDR"]
    data = [[Paragraph(str(h), sHDR) for h in headers]]
    for row in rows:
        data.append([Paragraph(str(c), sS) for c in row])
    t = Table(data, colWidths=col_widths)
    t.setStyle(TableStyle(pdf_base_tbl_style(len(rows), C)))
    return t


# ══════════════════════════════════════════════════════════════════════════════
# PDF — meta / banner helpers
# ══════════════════════════════════════════════════════════════════════════════

def pdf_meta_table(kv_rows: list, col_widths: list, C: dict, styles: dict):
    """
    Build a two-column key-value meta table (e.g. Project, Scanned At).

    kv_rows    : list of (label_str, value_str)
    col_widths : [label_width, value_width] in ReportLab units
    """
    from reportlab.platypus import Table, TableStyle, Paragraph

    sB, sN = styles["sB"], styles["sN"]
    data = [[Paragraph(k, sB), Paragraph(v, sN)] for k, v in kv_rows]
    t = Table(data, colWidths=col_widths)
    t.setStyle(TableStyle([
        ("BACKGROUND",    (0, 0), (0, -1), C["LBLU"]),
        ("GRID",          (0, 0), (-1, -1), 0.4, C["GRD"]),
        ("LEFTPADDING",   (0, 0), (-1, -1), 5),
        ("TOPPADDING",    (0, 0), (-1, -1), 3),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
    ]))
    return t


def pdf_gate_paragraph(gate_text: str, gate_hex: str, FB: str):
    """
    Return a centred Paragraph styled in the gate colour (PASS/WARN/FAIL banner).
    """
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER
    from reportlab.platypus import Paragraph

    style = ParagraphStyle(
        "sGate", fontSize=11, fontName=FB,
        alignment=TA_CENTER,
        textColor=colors.HexColor(gate_hex),
    )
    return Paragraph(gate_text, style)


# ══════════════════════════════════════════════════════════════════════════════
# DOCX — table helpers
# ══════════════════════════════════════════════════════════════════════════════

def _docx_set_font(run, font_name: str = "TH Sarabun New") -> None:
    """
    Set all four font slots on a python-docx run so Word uses the same font
    for Latin, East-Asian, complex-script, and high-ANSI characters.
    This prevents Word's theme font (Calibri/Cambria) from overriding the choice.
    """
    from docx.oxml.ns import qn
    from lxml import etree

    run.font.name = font_name
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = etree.SubElement(rPr, qn("w:rFonts"))
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rFonts.set(qn(attr), font_name)
    # Remove theme-font overrides so our explicit name wins
    for attr in ("w:asciiTheme", "w:hAnsiTheme", "w:eastAsiaTheme", "w:cstheme"):
        try:
            rFonts.attrib.pop(qn(attr))
        except KeyError:
            pass


def docx_set_bg(cell, hex_col: str) -> None:
    """Apply a solid background fill to a python-docx table cell."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_col.lstrip("#"))
    tcPr.append(shd)


def docx_set_col_width(table, col_index: int, width_cm: float) -> None:
    """Force an exact column width (in cm) on every cell in a column."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.shared import Cm

    twips = int(Cm(width_cm).twips)
    for row in table.rows:
        tc   = row.cells[col_index]._tc
        tcPr = tc.get_or_add_tcPr()
        tcW  = OxmlElement("w:tcW")
        tcW.set(qn("w:w"),    str(twips))
        tcW.set(qn("w:type"), "dxa")
        old = tcPr.find(qn("w:tcW"))
        if old is not None:
            tcPr.remove(old)
        tcPr.append(tcW)


def docx_fix_table_width(table, width_cm: float) -> None:
    """Set the overall table width to an exact value in cm."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.shared import Cm

    tbl  = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"),    str(int(Cm(width_cm).twips)))
    tblW.set(qn("w:type"), "dxa")
    old = tblPr.find(qn("w:tblW"))
    if old is not None:
        tblPr.remove(old)
    tblPr.append(tblW)


def docx_make_table(doc, headers: list, rows_data: list,
                    col_widths_cm: list, total_width_cm: float = 17.0,
                    hdr_bg: str = "1A5276", row_bg: str = "EBF5FB"):
    """
    Add and return a styled table to a python-docx Document.

    doc            : active Document instance
    headers        : list of header strings (pass [] to skip header row)
    rows_data      : list of lists of cell values
    col_widths_cm  : list of column widths in cm
    total_width_cm : table total width in cm (default A4 body = 17 cm)
    hdr_bg         : hex colour for header background
    row_bg         : hex colour for alternate data row background
    """
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    n_cols = len(headers) if headers else (len(rows_data[0]) if rows_data else 1)
    t = doc.add_table(rows=(1 if headers else 0) + len(rows_data), cols=n_cols)
    t.style = "Table Grid"
    docx_fix_table_width(t, total_width_cm)

    if headers:
        for i, h in enumerate(headers):
            c   = t.rows[0].cells[i]
            c.text = h
            run = c.paragraphs[0].runs[0]
            run.bold = True
            run.font.size = Pt(8)
            _docx_set_font(run, "TH Sarabun New")
            docx_set_bg(c, hdr_bg)
            run.font.color.rgb = RGBColor(255, 255, 255)
            c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

    for ri, row in enumerate(rows_data):
        off = 1 if headers else 0
        for ci, val in enumerate(row):
            c = t.rows[ri + off].cells[ci]
            c.text = str(val or "")
            for run in c.paragraphs[0].runs:
                run.font.size = Pt(8)
                _docx_set_font(run, "TH Sarabun New")
            c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            if ri % 2 == 0:
                docx_set_bg(c, row_bg)

    for ci, w in enumerate(col_widths_cm):
        docx_set_col_width(t, ci, w)

    return t


def docx_add_heading(doc, text: str, level: int = 1,
                     color_hex: str = "1A5276") -> None:
    """Add a heading paragraph with the brand colour and consistent font."""
    from docx.shared import RGBColor

    p = doc.add_heading(text, level=level)
    for r in p.runs:
        _docx_set_font(r, "TH Sarabun New")
        r.font.color.rgb = RGBColor.from_string(color_hex)


def docx_gate_paragraph(doc, gate_text: str, gate_hex: str) -> None:
    """Append a bold, coloured gate-status paragraph to the document."""
    from docx.shared import Pt, RGBColor

    p   = doc.add_paragraph()
    run = p.add_run(gate_text)
    run.bold           = True
    run.font.size      = Pt(12)
    _docx_set_font(run, "TH Sarabun New")
    run.font.color.rgb = RGBColor.from_string(gate_hex.lstrip("#"))


# ══════════════════════════════════════════════════════════════════════════════
# XLSX — cell helpers
# ══════════════════════════════════════════════════════════════════════════════

def xlsx_styles() -> dict:
    """
    Return a dict of pre-built openpyxl style objects shared across all sheets.

    Keys: BFILL, LFILL, WFILL, HFILL, BDR, WHDR, BBLUE, NRM, CTR, LFT.
    """
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

    th = Side(style="thin", color="AED6F1")
    return dict(
        BFILL = PatternFill("solid", fgColor="1A5276"),
        LFILL = PatternFill("solid", fgColor="EBF5FB"),
        WFILL = PatternFill("solid", fgColor="FDFEFE"),
        HFILL = PatternFill("solid", fgColor="D6EAF8"),
        BDR   = Border(left=th, right=th, top=th, bottom=th),
        WHDR  = Font(bold=True, color="FFFFFF",  name="Arial", size=9),
        BBLUE = Font(bold=True, color="1A5276",  name="Arial", size=9),
        NRM   = Font(name="Arial", size=9),
        CTR   = Alignment(horizontal="center", vertical="center", wrap_text=True),
        LFT   = Alignment(horizontal="left",   vertical="center", wrap_text=True),
    )


def xlsx_hcell(ws, row: int, col: int, value, S: dict, width: float = None):
    """Write a dark-blue header cell; optionally set column width."""
    from openpyxl.utils import get_column_letter

    cell            = ws.cell(row=row, column=col, value=value)
    cell.font       = S["WHDR"]
    cell.fill       = S["BFILL"]
    cell.alignment  = S["CTR"]
    cell.border     = S["BDR"]
    if width:
        ws.column_dimensions[get_column_letter(col)].width = width
    return cell


def xlsx_dcell(ws, row: int, col: int, value, S: dict,
               alt: bool = False, bold: bool = False,
               color: str = None, align=None):
    """Write a standard data cell with alternating row shading."""
    from openpyxl.styles import Font

    cell           = ws.cell(row=row, column=col, value=value)
    cell.fill      = S["LFILL"] if alt else S["WFILL"]
    cell.font      = Font(bold=bold, name="Arial", size=9, color=color or "000000")
    cell.alignment = align if align is not None else S["CTR"]
    cell.border    = S["BDR"]
    return cell


def xlsx_title_row(ws, value: str, S: dict, end_col: int = 6) -> None:
    """Merge A1:end_col1 and write a styled title row."""
    from openpyxl.utils import get_column_letter

    col_letter = get_column_letter(end_col)
    ws.merge_cells(f"A1:{col_letter}1")
    tc            = ws["A1"]
    tc.value      = value
    tc.font       = S["BBLUE"].__class__(bold=True, size=12, color="1A5276", name="Arial")
    tc.alignment  = S["CTR"]
    tc.fill       = S["HFILL"]
    ws.row_dimensions[1].height = 26


def xlsx_section_header(ws, row: int, label: str, S: dict, end_col: int = 6) -> None:
    """Merge a full-width section-header row with light-blue fill."""
    from openpyxl.utils import get_column_letter
    from openpyxl.styles import Font

    ws.merge_cells(f"A{row}:{get_column_letter(end_col)}{row}")
    cell           = ws.cell(row=row, column=1, value=label)
    cell.font      = Font(bold=True, color="1A5276", name="Arial", size=10)
    cell.fill      = S["HFILL"]
    cell.alignment = S["LFT"]
    cell.border    = S["BDR"]
    ws.row_dimensions[row].height = 18

def docx_color_cell_text(table, row_idx: int, col_idx: int, hex_col: str, bold: bool = True) -> None:
    from docx.shared import RGBColor
    cell = table.rows[row_idx].cells[col_idx]
    for p in cell.paragraphs:
        for run in p.runs:
            if hex_col:
                run.font.color.rgb = RGBColor.from_string(hex_col.lstrip("#"))
            if bold:
                run.bold = True

def docx_set_default_font(doc, font_name: str = "Arial") -> None:
    from docx.oxml.ns import qn
    from lxml import etree
    
    style = doc.styles['Normal']
    style.font.name = font_name
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = etree.SubElement(rPr, qn("w:rFonts"))
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rFonts.set(qn(attr), font_name)