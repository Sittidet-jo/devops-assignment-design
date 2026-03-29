#!/usr/bin/env python3
"""
DAST (OWASP ZAP) Report Generator
PDF + DOCX + XLSX from zap_report.json
Gate: HIGH > 0 → FAIL, MEDIUM > 0 → WARN, else → PASS
Bilingual: Thai + English
"""
import json, sys, os, re
from datetime import datetime
from collections import defaultdict

from report_utils import (
    load_font, make_pdf_styles, pdf_colors, pdf_base_tbl_style,
    make_pdf_table, pdf_meta_table, pdf_gate_paragraph,
    docx_make_table, docx_add_heading, docx_gate_paragraph,docx_color_cell_text, docx_set_default_font,
    xlsx_styles, xlsx_hcell, xlsx_dcell, xlsx_title_row, xlsx_section_header,
)

def clean_html(text: str) -> str:
    if not text: return ""
    return re.sub(r'<[^>]+>', '', text).replace('&quot;', '"').strip()

def risk_color(risk: str) -> str:
    return {"3": "#C0392B", "2": "#E67E22", "1": "#D4AC0D", "0": "#2E86C1"}.get(str(risk), "#7F8C8D")

def risk_label(risk: str) -> str:
    return {"3": "High", "2": "Medium", "1": "Low", "0": "Informational"}.get(str(risk), "Unknown")

def parse_report(report: dict) -> dict:
    site_list = report.get("site", [])
    if not site_list:
        return {
            "total": 0, "target": "", "scanned_at": datetime.now().strftime("%Y-%m-%d %H:%M"),
            "counts": {}, "high_med": [], "low_info": [], "all_alerts": [], "gate": "PASS",
        }
    all_alerts = []; counts = defaultdict(int)
    for site in site_list:
        host = site.get("@host", "")
        for alert in site.get("alerts", []):
            risk = alert.get("riskcode", "0"); counts[risk] += 1
            all_alerts.append({
                "name": alert.get("alert", ""), "risk": risk,
                "risk_desc": alert.get("riskdesc", ""), 
                "desc": clean_html(alert.get("desc", "")),
                "solution": clean_html(alert.get("solution", "")),
                "count": alert.get("count", "1"),
                "host": host, "instances": alert.get("instances", []),
            })
    all_alerts.sort(key=lambda x: x["risk"], reverse=True)
    high = int(counts.get("3", 0)); med = int(counts.get("2", 0))
    gate = "FAIL" if high > 0 else ("WARN" if med > 0 else "PASS")
    return {
        "total": len(all_alerts), "target": site_list[0].get("@name", ""),
        "scanned_at": datetime.now().strftime("%Y-%m-%d %H:%M"),
        "counts": dict(counts), "gate": gate,
        "high_med": [a for a in all_alerts if a["risk"] in ("3", "2")],
        "low_info":  [a for a in all_alerts if a["risk"] in ("1", "0")],
        "all_alerts": all_alerts,
    }

TH = {
    "title":      "รายงานการทดสอบความปลอดภัยเว็บ (DAST) / Web Security Test Report",
    "summary":    "สรุปผลการสแกน / Scan Summary",
    "high_med":   "High & Medium — รายการ / All",
    "low_info":   "Low & Informational — รายการ / All",
    "no_alert":   "ไม่พบ Alert / No alerts found",
    "gate_pass":  "✅ PASS — ไม่พบ High/Medium Risk / No High or Medium Risks",
    "gate_warn":  "⚠️ WARN — พบ Medium Risk (Pipeline: UNSTABLE) / Medium Risks Found",
    "gate_fail":  "❌ FAIL — พบ High Risk (Pipeline: ABORTED) / High Risks Found",
    "target":     "เป้าหมาย / Target",
    "scanned":    "สแกนเวลา / Scanned At",
    "threshold":  "เกณฑ์การตัดสิน / Threshold",
    "repo":       "GitLab Repository",
    "branch":     "Branch",
    "repo_url":   "Repository URL",
    "risk":       "Risk Level",
    "count":      "Count",
    "alert_name": "Alert",
    "risk_col":   "Risk",
    "instances":  "Instances",
    "desc_col":   "Description",
    "solution":   "Solution",
}


GATE_COLOR = {"PASS": "#1E8449", "WARN": "#E67E22", "FAIL": "#C0392B"}
GATE_TEXT  = {k: TH[f"gate_{k.lower()}"] for k in GATE_COLOR}

def generate_pdf(data: dict, out_path: str) -> None:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.lib.units import mm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable, PageBreak

    FN, FB = load_font()
    ST = make_pdf_styles(FN, FB); C  = pdf_colors()
    sT, sH1, sN, sS, sB, sHDR = ST["sT"], ST["sH1"], ST["sN"], ST["sS"], ST["sB"], ST["sHDR"]

    W, _ = A4; CW = W - 30 * mm
    doc = SimpleDocTemplate(out_path, pagesize=A4, leftMargin=15*mm, rightMargin=15*mm, topMargin=18*mm, bottomMargin=18*mm)

    counts = data["counts"]; story = []

    story.append(Paragraph(TH["title"], sT))
    story.append(HRFlowable(width="100%", thickness=2, color=C["BLUE"]))
    story.append(Spacer(1, 3*mm))

    repo     = os.environ.get("REPO_PATH", "-")
    branch   = os.environ.get("BRANCH_NAME", "-")
    repo_url = os.environ.get("REPO_URL", "-")

    kv = [
        (TH["scanned"], data["scanned_at"]), 
        (TH["target"], data["target"]),
        (TH["repo"], repo),
        (TH["branch"], branch),
        (TH["repo_url"], repo_url)
    ]
    if data.get("criteria"):
        kv.append((TH["threshold"], data["criteria"]))
    story.append(pdf_meta_table(kv, [40*mm, CW - 40*mm], C, ST))
    story.append(Spacer(1, 3*mm))

    gate = data["gate"]
    story.append(pdf_gate_paragraph(GATE_TEXT.get(gate, ""), GATE_COLOR.get(gate, "#7F8C8D"), FB))
    story.append(Spacer(1, 4*mm))

    story.append(Paragraph(TH["summary"], sH1))
    sev_rows = [(risk_label(rc), str(counts.get(rc, 0)), rc) for rc in ["3", "2", "1", "0"] if counts.get(rc, 0) > 0]
    sum_style = pdf_base_tbl_style(len(sev_rows), C)
    for i, (_, _, rc) in enumerate(sev_rows):
        c = colors.HexColor(risk_color(rc))
        sum_style += [("TEXTCOLOR", (0, i+1), (0, i+1), c), ("FONTNAME", (0, i+1), (0, i+1), FB)]
    st = Table(
        [[Paragraph(TH["risk"], sHDR), Paragraph(TH["count"], sHDR)]] +
        [[Paragraph(f'<font color="{risk_color(rc)}"><b>{r}</b></font>', sS), Paragraph(c, sS)] for r, c, rc in sev_rows],
        colWidths=[100*mm, CW - 100*mm],
    )
    st.setStyle(TableStyle(sum_style)); story.append(st); story.append(Spacer(1, 4*mm))

    def alert_table(alert_list: list, label: str) -> None:
        story.append(Paragraph(f"<b>{label}</b> ({len(alert_list)} items)", sN))
        story.append(Spacer(1, 1*mm))
        if not alert_list:
            story.append(Paragraph(TH["no_alert"], sS)); story.append(Spacer(1, 3*mm)); return
        hdrs = [TH["alert_name"], TH["risk_col"], TH["instances"], TH["desc_col"]]
        cws  = [55*mm, 25*mm, 18*mm, CW - 98*mm]
        a_style = pdf_base_tbl_style(len(alert_list), C)
        for i, a in enumerate(alert_list):
            c = colors.HexColor(risk_color(a["risk"]))
            a_style += [("TEXTCOLOR", (1, i+1), (1, i+1), c), ("FONTNAME", (1, i+1), (1, i+1), FB)]
        rows = [[a["name"], f'<font color="{risk_color(a["risk"])}"><b>{risk_label(a["risk"])}</b></font>', a["count"],
                 a["desc"][:150] + "…" if len(a["desc"]) > 150 else a["desc"]] for a in alert_list]
        t = Table(
            [[Paragraph(h, sHDR) for h in hdrs]] +
            [[Paragraph(str(c), sS) for c in row] for row in rows],
            colWidths=cws,
        )
        t.setStyle(TableStyle(a_style)); story.append(t); story.append(Spacer(1, 4*mm))

    story.append(PageBreak())
    alert_table(data["high_med"], TH["high_med"])
    story.append(PageBreak())
    alert_table(data["low_info"], TH["low_info"])
    doc.build(story)

def generate_docx(data: dict, out_path: str) -> None:
    from docx import Document
    from docx.shared import Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import RGBColor

    TW = 17.0
    doc = Document()
    docx_set_default_font(doc, "Arial")
    for sec in doc.sections:
        sec.top_margin = sec.bottom_margin = Cm(2)
        sec.left_margin = sec.right_margin = Cm(2)

    counts = data["counts"]

    t = doc.add_heading(TH["title"], 0); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in t.runs: r.font.color.rgb = RGBColor.from_string("1A5276")

    repo     = os.environ.get("REPO_PATH", "-")
    branch   = os.environ.get("BRANCH_NAME", "-")
    repo_url = os.environ.get("REPO_URL", "-")

    kv = [
        [TH["scanned"], data["scanned_at"]], 
        [TH["target"], data["target"]],
        [TH["repo"], repo],
        [TH["branch"], branch],
        [TH["repo_url"], repo_url]
    ]
    if data.get("criteria"): kv.append([TH["threshold"], data["criteria"]])
    docx_make_table(doc, [], kv, [4.5, TW - 4.5])
    doc.add_paragraph()

    gate = data["gate"]
    docx_gate_paragraph(doc, GATE_TEXT.get(gate, ""), GATE_COLOR.get(gate, "#7F8C8D"))
    doc.add_paragraph()

    docx_add_heading(doc, TH["summary"])
    sev_rows = [(rc, risk_label(rc), str(counts.get(rc, 0))) for rc in ["3", "2", "1", "0"] if counts.get(rc, 0) > 0]
    t_sum = docx_make_table(doc, [TH["risk"], TH["count"]], [[rl, c] for rc, rl, c in sev_rows], [8.0, TW - 8.0])
    for i, (rc, rl, c) in enumerate(sev_rows):
        docx_color_cell_text(t_sum, i + 1, 0, risk_color(rc))
    doc.add_paragraph()

    def add_alert_table(alert_list: list, label: str) -> None:
        docx_add_heading(doc, label, level=2)
        if not alert_list:
            doc.add_paragraph(TH["no_alert"]); return
        t_a = docx_make_table(
            doc,
            [TH["alert_name"], TH["risk_col"], TH["instances"], TH["desc_col"]],
            [[a["name"], risk_label(a["risk"]), a["count"], a["desc"][:200]] for a in alert_list],
            [6.0, 2.5, 2.0, TW - 10.5],
        )
        for i, a in enumerate(alert_list):
            docx_color_cell_text(t_a, i + 1, 1, risk_color(a["risk"]))
        doc.add_paragraph()

    doc.add_page_break()
    add_alert_table(data["high_med"], TH["high_med"])
    add_alert_table(data["low_info"], TH["low_info"])
    doc.save(out_path)

def generate_xlsx(data: dict, out_path: str) -> None:
    from openpyxl import Workbook
    from openpyxl.styles import Font

    wb = Workbook(); counts = data["counts"]
    S  = xlsx_styles()

    ws1 = wb.active; ws1.title = "📊 Summary"
    xlsx_title_row(ws1, TH["title"], S, end_col=4)

    repo     = os.environ.get("REPO_PATH", "-")
    branch   = os.environ.get("BRANCH_NAME", "-")
    repo_url = os.environ.get("REPO_URL", "-")

    _meta = [
        (TH["scanned"], data["scanned_at"]),
        (TH["target"], data["target"]),
        (TH["repo"], repo),
        (TH["branch"], branch),
        (TH["repo_url"], repo_url)
    ]
    if data.get("criteria"):
        _meta.append((TH["threshold"], data["criteria"]))

    for i, (k, v) in enumerate(_meta, start=2):
        c1 = ws1.cell(row=i, column=1, value=k); c1.font = S["BBLUE"]; c1.fill = S["LFILL"]; c1.border = S["BDR"]; c1.alignment = S["LFT"]
        ws1.merge_cells(f"B{i}:D{i}"); c2 = ws1.cell(row=i, column=2, value=v); c2.font = S["NRM"]; c2.border = S["BDR"]; c2.alignment = S["LFT"]

    row = 2 + len(_meta) + 1; gate = data["gate"]; gate_hex = GATE_COLOR.get(gate, "7F8C8D").lstrip("#")
    ws1.merge_cells(f"A{row}:D{row}"); gc = ws1.cell(row=row, column=1, value=GATE_TEXT.get(gate, ""))
    gc.font = Font(bold=True, color=gate_hex, name="Arial", size=11); gc.alignment = S["CTR"]; gc.border = S["BDR"]
    gc.fill = S["WFILL"]; ws1.row_dimensions[row].height = 22; row += 2

    xlsx_section_header(ws1, row, TH["summary"], S, end_col=4); row += 1
    xlsx_hcell(ws1, row, 1, TH["risk"], S, 25); xlsx_hcell(ws1, row, 2, TH["count"], S, 15); row += 1
    for i, rc in enumerate(["3", "2", "1", "0"]):
        cnt = counts.get(rc, 0)
        if cnt == 0: continue
        alt = i % 2 == 0; sc = risk_color(rc).lstrip("#")
        xlsx_dcell(ws1, row, 1, risk_label(rc), S, alt, bold=True, color=sc)
        xlsx_dcell(ws1, row, 2, cnt, S, alt); row += 1

    ws2 = wb.create_sheet("🔴 Alerts Detail")
    xlsx_title_row(ws2, f"Alerts Detail ({data['total']} items)", S, end_col=5)
    for col, (h, w) in enumerate(zip([TH["alert_name"], TH["risk_col"], TH["instances"], TH["desc_col"], TH["solution"]], [45, 15, 12, 50, 50]), 1):
        xlsx_hcell(ws2, 2, col, h, S, w)

    if not data["all_alerts"]:
        ws2.cell(row=3, column=1, value=TH["no_alert"]).font = S["NRM"]
    else:
        for i, a in enumerate(data["all_alerts"]):
            r = i + 3; alt = i % 2 == 0; sc = risk_color(a["risk"]).lstrip("#")
            xlsx_dcell(ws2, r, 1, a["name"],              S, alt, align=S["LFT"])
            xlsx_dcell(ws2, r, 2, risk_label(a["risk"]),  S, alt, bold=True, color=sc)
            xlsx_dcell(ws2, r, 3, a["count"],             S, alt)
            xlsx_dcell(ws2, r, 4, a["desc"],              S, alt, align=S["LFT"])
            xlsx_dcell(ws2, r, 5, a["solution"],          S, alt, align=S["LFT"])
            ws2.row_dimensions[r].height = 45

            ws2.column_dimensions["A"].width = 45   # alert name
            ws2.column_dimensions["B"].width = 16   # risk
            ws2.column_dimensions["C"].width = 12   # instances
            ws2.column_dimensions["D"].width = 55   # description
            ws2.column_dimensions["E"].width = 55   # solution

    wb.save(out_path)

def main():
    json_path = sys.argv[1] if len(sys.argv) > 1 else "zap_report.json"
    out_dir   = sys.argv[2] if len(sys.argv) > 2 else os.path.dirname(json_path) or "reports"
    criteria  = sys.argv[3] if len(sys.argv) > 3 else None

    if not os.path.exists(json_path):
        print(f"[ERROR] File not found: {json_path}", file=sys.stderr); sys.exit(1)
    os.makedirs(out_dir, exist_ok=True)

    report = json.load(open(json_path, encoding="utf-8"))
    data   = parse_report(report)
    if criteria: data["criteria"] = criteria

    project_name = os.environ.get("PROJECT_NAME", "")
    deploy_env   = os.environ.get("DEPLOY_ENV", "")
    date = datetime.now().strftime("%Y%m%d")

    if project_name and deploy_env:
        base = f"zap-{project_name}-{deploy_env}-{date}"
    else:
        base = f"zap-report_{date}"
    print(f"[INFO] DAST: target={data['target']}  gate={data['gate']}")
    generate_pdf (data, os.path.join(out_dir, f"{base}.pdf"))
    generate_docx(data, os.path.join(out_dir, f"{base}.docx"))
    generate_xlsx(data, os.path.join(out_dir, f"{base}.xlsx"))

if __name__ == "__main__":
    main()